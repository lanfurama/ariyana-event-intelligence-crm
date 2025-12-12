#!/usr/bin/env python3
"""Script to read and extract text from Word document (.docx)"""

import sys
from docx import Document

def read_docx(file_path):
    """Read and extract text from a Word document"""
    try:
        doc = Document(file_path)
        
        print("=" * 80)
        print(f"File: {file_path}")
        print("=" * 80)
        print()
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            if table_data:
                tables.append(table_data)
        
        # Print paragraphs
        if paragraphs:
            print("PARAGRAPHS:")
            print("-" * 80)
            for i, para in enumerate(paragraphs, 1):
                print(f"{i}. {para}")
            print()
        
        # Print tables
        if tables:
            print("TABLES:")
            print("-" * 80)
            for table_idx, table in enumerate(tables, 1):
                print(f"\nTable {table_idx}:")
                for row in table:
                    print(" | ".join(row))
                print()
        
        # Summary
        print("=" * 80)
        print(f"Summary: {len(paragraphs)} paragraphs, {len(tables)} tables")
        print("=" * 80)
        
        return {
            'paragraphs': paragraphs,
            'tables': tables,
            'text': '\n\n'.join(paragraphs)
        }
        
    except Exception as e:
        print(f"Error reading file: {e}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    read_docx(file_path)

